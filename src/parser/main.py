# read docx file into headers
from dataclasses import dataclass
from pathlib import Path
from typing import List

from docx import Document
from openpyxl import Workbook


@dataclass
class WordDocument:
    path: Path

    @classmethod
    def make(cls, path: Path) -> "WordDocument":
        return WordDocument(path=path)


@dataclass
class Header:
    level: int
    text: str
    children: List["Header"] | None = None

    def add_child(self, child: "Header"):
        if self.children is None:
            self.children = []
        self.children.append(child)

    def compare(self, other: "Header") -> List[str]:
        missing_items = []

        # Check if the current header is missing in other instance
        if self.level != other.level or self.text != other.text:
            missing_items.append(self.text)

        # Recursively check for missing children in other instance
        if self.children is not None:
            for child in self.children:
                matching_child = next(
                    (
                        c
                        for c in other.children or []
                        if c.level == child.level and c.text == child.text
                    ),
                    None,
                )
                if matching_child is None:
                    missing_items.append(child.text)
                else:
                    missing_items.extend(child.compare(matching_child))

        return missing_items


def build_header_hierarchy(headers: List[Header]) -> Header:
    """_summary_

    Args:
        headers (List[Header]): _description_

    Returns:
        Header: _description_
    """
    root_header = Header(0, "ROOT")
    level_1_header = None
    level_2_header = None
    for header in reversed(headers):
        if header.level == 1:
            root_header.add_child(header)
            level_1_header = header
            level_2_header = None
        elif header.level == 2:
            level_1_header.add_child(header)
            level_2_header = header
        elif header.level == 3:
            if level_2_header is None:
                level_2_header = Header(2, "")
                level_1_header.add_child(level_2_header)
            level_2_header.add_child(header)
    return root_header


def flatten_root(root_header: Header) -> list[str]:
    headers_list = []
    for child in root_header.children:
        print(child.text)
        headers_list.append((child.level, child.text))
        if child.children:
            for childchild in child.children:
                print(childchild.text)
                headers_list.append((childchild.level, childchild.text))
                if childchild.children:
                    for childchildchild in childchild.children:
                        print(childchildchild.text)
                        headers_list.append(
                            (childchildchild.level, childchildchild.text)
                        )
    return headers_list


def load_headers_from_file(file: Path) -> List[Header]:
    project_document = WordDocument.make(file)
    word_document = Document(project_document.path)
    content = word_document.paragraphs
    headers = [para for para in content if para.style.name.startswith("Heading")]
    list_of_headers = []
    for i, header in enumerate(reversed(headers)):
        header_type = header.style.name
        match header_type:
            case "Heading 1":
                level = 1
            case "Heading 2":
                level = 2
            case "Heading 3":
                level = 3
            case _:
                level = 0
        list_of_headers.append(Header(level=level, text=header.text))
    return list_of_headers


def write_to_excel_file(root_header: Header):
    wb = Workbook()
    ws = wb.active
    ws.append(["Level", "Heading"])
    for head in flatten_root(root_header=root_header):
        ws.append([head[0], head[1]])
    wb.save("output.xlsx")


if __name__ == "__main__":
    input_file = Path("tests", "fixtures", "parser_test.docx")
    list_of_headers = load_headers_from_file(input_file)
    root_header = build_header_hierarchy(list_of_headers)

    check_file = Path("tests", "fixtures", "parser_test_no_match.docx")
    list_of_headers_to_check = load_headers_from_file(check_file)
    check_header = build_header_hierarchy(list_of_headers_to_check)

    write_to_excel_file(root_header)

    workbook = Workbook()
    sheet = workbook.active

    max_col = sheet.max_column
    sheet.cell(row=1, column=1, value="Level")
    sheet.cell(row=1, column=2, value="Heading")

    for row, value in enumerate(flatten_root(root_header=root_header), start=2):
        sheet.cell(row=row, column=max_col, value=value[0])
        sheet.cell(row=row, column=max_col + 1, value=value[1])

    max_col = sheet.max_column + 2
    sheet.cell(row=1, column=4, value="Level")
    sheet.cell(row=1, column=5, value="Heading")
    for row, value in enumerate(flatten_root(root_header=check_header), start=2):
        sheet.cell(row=row, column=max_col, value=value[0])
        sheet.cell(row=row, column=max_col + 1, value=value[1])

    sheet.cell(row=sheet.max_row + 2, column=1, value="Missing columns")
    sheet.append(root_header.compare(check_header))

    workbook.save("output2.xlsx")
